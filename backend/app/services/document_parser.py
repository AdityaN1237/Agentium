"""
Document Parser Service.
Handles text extraction from PDFs, DOCX, TXT, and images (OCR).
Supports single files, folders, and ZIP archives.
"""
import os
import io
import zipfile
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Union, Optional
import logging
import base64
import asyncio
from app.services.local_vision import get_local_vision

logger = logging.getLogger(__name__)

# PDF extraction
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF not installed. PDF extraction will be limited.")

# DOCX extraction
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed. DOCX extraction unavailable.")

# OCR for images
try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    logger.warning("pytesseract/PIL not installed. OCR unavailable.")


class DocumentParser:
    """
    Unified document parser supporting multiple formats.
    
    Supported formats:
    - PDF (via PyMuPDF)
    - DOCX (via python-docx)
    - TXT (direct read)
    - Images (via pytesseract OCR)
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg', '.tiff', '.bmp'}
    
    def __init__(self):
        self.stats = {
            "total_processed": 0,
            "successful": 0,
            "failed": 0
        }
    
    async def extract_text(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from a file based on its extension (Async).
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            return await self.extract_from_bytes(content, file_path.name)
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            self.stats["failed"] += 1
            return ""
    
    async def extract_from_bytes(self, content: bytes, filename: str) -> str:
        """
        Extract text from bytes content (Async).
        
        Args:
            content: File content as bytes
            filename: Original filename for extension detection
            
        Returns:
            Extracted text content
        """
        extension = Path(filename).suffix.lower()
        
        try:
            if extension == '.pdf':
                return await self._extract_pdf_bytes(content)
            elif extension in {'.docx', '.doc'}:
                return self._extract_docx_bytes(content)
            elif extension == '.txt':
                return content.decode('utf-8', errors='ignore')
            elif extension in {'.png', '.jpg', '.jpeg', '.tiff', '.bmp'}:
                return self._extract_image_ocr_bytes(content)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return ""
        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {e}")
            return ""
    
    async def process_folder(self, folder_path: Union[str, Path]) -> List[Dict[str, Any]]:
        """
        Process all supported documents in a folder (Async).
        
        Args:
            folder_path: Path to the folder
            
        Returns:
            List of dicts with filename and extracted text
        """
        folder_path = Path(folder_path)
        results = []
        
        if not folder_path.exists():
            logger.error(f"Folder not found: {folder_path}")
            return results
        
        for file_path in folder_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                text = await self.extract_text(file_path)
                if text:
                    results.append({
                        "filename": file_path.name,
                        "path": str(file_path),
                        "text": text,
                        "size_bytes": file_path.stat().st_size
                    })
                    self.stats["successful"] += 1
                self.stats["total_processed"] += 1
        
        logger.info(f"Processed {len(results)} documents from {folder_path}")
        return results
    
    async def process_zip(self, zip_path: Union[str, Path]) -> List[Dict[str, Any]]:
        """
        Extract and process all supported documents from a ZIP file (Async).
        
        Args:
            zip_path: Path to the ZIP file
            
        Returns:
            List of dicts with filename and extracted text
        """
        zip_path = Path(zip_path)
        results = []
        
        with zipfile.ZipFile(zip_path, 'r') as zf:
            for file_info in zf.infolist():
                if file_info.is_dir():
                    continue
                    
                filename = Path(file_info.filename).name
                extension = Path(filename).suffix.lower()
                
                if extension in self.SUPPORTED_EXTENSIONS:
                    content = zf.read(file_info.filename)
                    text = await self.extract_from_bytes(content, filename)
                    
                    if text:
                        results.append({
                            "filename": filename,
                            "path": file_info.filename,
                            "text": text,
                            "size_bytes": file_info.file_size
                        })
                        self.stats["successful"] += 1
                    self.stats["total_processed"] += 1
        
        logger.info(f"Processed {len(results)} documents from ZIP: {zip_path}")
        return results
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF."""
        if not HAS_PYMUPDF:
            logger.error("PyMuPDF not available for PDF extraction")
            return ""
        
        text = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text.append(page.get_text())
        
        return "\n".join(text)
    
    async def _extract_pdf_bytes(self, content: bytes) -> str:
        """
        Extract text AND images from PDF bytes using PyMuPDF + Local Vision (Florence-2).
        """
        if not HAS_PYMUPDF:
            return ""
        
        full_text = []
        vision_service = get_local_vision()

        with fitz.open(stream=content, filetype="pdf") as doc:
            for page_num, page in enumerate(doc):
                # 1. Extract Text
                page_text = page.get_text()
                full_text.append(page_text)
                
                # 2. Extract Images
                try:
                    image_list = page.get_images(full=True)
                except Exception:
                    image_list = []
                    
                if not image_list:
                    continue

                logger.info(f"Found {len(image_list)} images on page {page_num+1}")
                
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        
                        # Handle CMYK or other non-RGB
                        if pix.n - pix.alpha > 3:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                            
                        # Convert to PNG bytes
                        img_bytes = pix.tobytes("png")
                        
                        # Run Analysis (Blocking, so run in thread)
                        logger.info(f"Analyzing image {img_index+1} on page {page_num+1} using Florence-2...")
                        
                        # Florence-2 inference
                        analysis_result = await asyncio.to_thread(
                            vision_service.analyze_image, 
                            img_bytes
                        )
                        
                        # Append description to text
                        full_text.append(f"\n[IMAGE_ANALYSIS_PAGE_{page_num+1}_IMG_{img_index+1}]\n{analysis_result}\n[END_IMAGE_ANALYSIS]\n")
                        
                    except Exception as e:
                        logger.error(f"Failed to process image {img_index} on page {page_num+1}: {e}")
                        continue
        
                # 3. Full Page Visual Scan (Fallback for Scanned/Handwritten pages)
                # If text is sparse (< 50 chars), it's likely a scan or image-heavy page.
                if len(page_text.strip()) < 50:
                    logger.info(f"Page {page_num+1} seems sparse/scanned using Fallback Vision Analysis...")
                    try:
                        # Render full page as high-res image
                        zoom = 2  # 2x zoom for better OCR
                        mat = fitz.Matrix(zoom, zoom)
                        pix = page.get_pixmap(matrix=mat)
                        
                        if pix.n - pix.alpha > 3:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                            
                        page_bytes = pix.tobytes("png")
                        
                        # Full page analysis
                        page_analysis = await asyncio.to_thread(
                            vision_service.analyze_image,
                            page_bytes
                        )
                        
                        full_text.append(f"\n[FULL_PAGE_VISUAL_ANALYSIS_PAGE_{page_num+1}]\n{page_analysis}\n[END_VISUAL_ANALYSIS]\n")
                        
                    except Exception as e:
                        logger.error(f"Full page analysis failed for page {page_num+1}: {e}")
        
        return "\n".join(full_text)
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        if not HAS_DOCX:
            logger.error("python-docx not available for DOCX extraction")
            return ""
        
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def _extract_docx_bytes(self, content: bytes) -> str:
        """Extract text from DOCX bytes."""
        if not HAS_DOCX:
            return ""
        
        doc = DocxDocument(io.BytesIO(content))
        return "\n".join([para.text for para in doc.paragraphs])
    
    def _extract_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _extract_image_ocr(self, file_path: Path) -> str:
        """Extract text from image using OCR."""
        if not HAS_OCR:
            logger.error("OCR dependencies not available")
            return ""
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {e}")
            return ""
    
    def _extract_image_ocr_bytes(self, content: bytes) -> str:
        """Extract text from image bytes using OCR."""
        if not HAS_OCR:
            return ""
        
        try:
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return ""
    
    def get_stats(self) -> Dict[str, int]:
        """Get processing statistics."""
        return self.stats.copy()
    
    def reset_stats(self):
        """Reset processing statistics."""
        self.stats = {"total_processed": 0, "successful": 0, "failed": 0}


# Singleton
_parser: Optional[DocumentParser] = None

def get_document_parser() -> DocumentParser:
    """Get singleton document parser instance."""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser
