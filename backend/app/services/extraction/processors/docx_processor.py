import os
import logging
from .base import BaseProcessor
from ..result import ConversionResult
from ..exceptions import ConversionError, FileNotFoundError

logger = logging.getLogger(__name__)

class DOCXProcessor(BaseProcessor):
    def can_process(self, file_path: str) -> bool:
        if not os.path.exists(file_path):
            logger.debug(f"File not found for DOCXProcessor: {file_path}")
            return False
        _, ext = os.path.splitext(str(file_path).lower())
        return ext in ['.docx', '.doc']

    def process(self, file_path: str) -> ConversionResult:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
        _, ext = os.path.splitext(str(file_path).lower())
        if ext == '.doc':
            logger.info(f"Processing .doc file: {file_path}")
            return self._process_doc_file(file_path)
        else:
            logger.info(f"Processing .docx file: {file_path}")
            return self._process_docx_file(file_path)

    def _process_doc_file(self, file_path: str) -> ConversionResult:
        try:
            import pypandoc
            content = pypandoc.convert_file(file_path, 'markdown')
            metadata = self.get_metadata(file_path)
            metadata.update({"file_type": "doc", "extractor": "pypandoc"})
            content = self._clean_content(content)
            logger.info(f"Successfully processed .doc file: {file_path}")
            return ConversionResult(content, metadata)
        except ImportError:
            logger.error("pypandoc is required for .doc file processing. Install it with: pip install pypandoc")
            raise ConversionError("pypandoc is required for .doc file processing. Install it with: pip install pypandoc")
        except Exception as e:
            logger.error(f"Failed to process .doc file {file_path}: {str(e)}")
            raise ConversionError(f"Failed to process .doc file {file_path}: {str(e)}")

    def _process_docx_file(self, file_path: str) -> ConversionResult:
        try:
            from docx import Document
            content_parts = []
            doc = Document(file_path)
            metadata = self.get_metadata(file_path)
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections),
                "file_type": "docx",
                "extractor": "python-docx"
            })
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.replace('Heading ', '')
                        try:
                            level_num = int(level)
                            content_parts.append(f"\n{'#' * min(level_num, 6)} {paragraph.text}\n")
                        except ValueError:
                            content_parts.append(f"\n## {paragraph.text}\n")
                    else:
                        content_parts.append(paragraph.text)
            for table_idx, table in enumerate(doc.tables):
                if self.preserve_layout:
                    content_parts.append(f"\n### Table {table_idx+1}\n")
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    header = table_data[0]
                    separator = ["---"] * len(header)
                    content_parts.append("| " + " | ".join(header) + " |")
                    content_parts.append("| " + " | ".join(separator) + " |")
                    for row in table_data[1:]:
                        content_parts.append("| " + " | ".join(row) + " |")
                    content_parts.append("")
            content = '\n'.join(content_parts)
            content = self._clean_content(content)
            logger.info(f"Successfully processed .docx file: {file_path}")
            return ConversionResult(content, metadata)
        except ImportError:
            logger.error("python-docx is required for .docx file processing. Install it with: pip install python-docx")
            raise ConversionError("python-docx is required for .docx file processing. Install it with: pip install python-docx")
        except Exception as e:
            logger.error(f"Failed to process .docx file {file_path}: {str(e)}")
            raise ConversionError(f"Failed to process .docx file {file_path}: {str(e)}")

    def _clean_content(self, content: str) -> str:
        lines = content.split('\n')
        cleaned_lines = []
        for line in lines:
            line = ' '.join(line.split())
            if line.strip():
                cleaned_lines.append(line)
        content = '\n'.join(cleaned_lines)
        content = content.replace('## ', '\n## ')
        content = content.replace('### ', '\n### ')
        return content.strip()
